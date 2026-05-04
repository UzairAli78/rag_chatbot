"""
Document Processor Module
──────────────────────────
Handles loading, parsing, and chunking of multiple document types:
  PDF, DOCX, TXT, Markdown, CSV

Uses LangChain document loaders and RecursiveCharacterTextSplitter
for smart, overlap-aware chunking.
"""

import os
import logging
from pathlib import Path
from typing import List

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from src.config import CHUNK_SIZE, CHUNK_OVERLAP, ALLOWED_EXTENSIONS

logger = logging.getLogger(__name__)


# ─── Standalone Helpers ──────────────────────────────────────────────────────

def extract_docx_text(path: str) -> str:
    """
    Extract plain text from a .docx file.
    Tries docx2txt first (fast), falls back to python-docx.
    """
    try:
        import docx2txt
        text = docx2txt.process(path)
        return text or ""
    except Exception:
        pass

    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Failed to extract DOCX text from {path}: {e}")
        return ""


# ─── Document Processor ──────────────────────────────────────────────────────

class DocumentProcessor:
    """
    Loads documents from disk and splits them into overlapping chunks
    suitable for embedding and retrieval.
    """

    def __init__(self) -> None:
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        )

    # ── Loaders ──────────────────────────────────────────────────────────────

    def _load_pdf(self, path: str) -> List[Document]:
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(path)
        return loader.load()

    def _load_docx(self, path: str) -> List[Document]:
        try:
            from langchain_community.document_loaders import Docx2txtLoader
            loader = Docx2txtLoader(path)
            docs = loader.load()
            if docs:
                return docs
        except Exception:
            pass
        # Fallback: extract text manually
        text = extract_docx_text(path)
        return [Document(page_content=text, metadata={})] if text else []

    def _load_txt(self, path: str) -> List[Document]:
        from langchain_community.document_loaders import TextLoader
        loader = TextLoader(path, encoding="utf-8", autodetect_encoding=True)
        return loader.load()

    def _load_md(self, path: str) -> List[Document]:
        # Try UnstructuredMarkdownLoader first; fall back to TextLoader
        try:
            from langchain_community.document_loaders import UnstructuredMarkdownLoader
            loader = UnstructuredMarkdownLoader(path)
            return loader.load()
        except Exception:
            return self._load_txt(path)

    def _load_csv(self, path: str) -> List[Document]:
        from langchain_community.document_loaders import CSVLoader
        loader = CSVLoader(path, encoding="utf-8")
        return loader.load()

    # ── Public API ───────────────────────────────────────────────────────────

    def load_document(self, file_path: str) -> List[Document]:
        """
        Load a document by detecting its extension and using the correct loader.

        Args:
            file_path: Absolute path to the file.

        Returns:
            List of LangChain Document objects (one per page/row/section).

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the extension is not supported.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{ext}'. "
                f"Supported: {', '.join(ALLOWED_EXTENSIONS)}"
            )

        loader_map = {
            ".pdf":  self._load_pdf,
            ".docx": self._load_docx,
            ".txt":  self._load_txt,
            ".md":   self._load_md,
            ".csv":  self._load_csv,
        }

        loader_fn = loader_map[ext]

        try:
            docs = loader_fn(file_path)
        except Exception as e:
            logger.error(f"Error loading {file_path}: {e}")
            raise RuntimeError(f"Could not load '{path.name}': {e}") from e

        # Attach source metadata to every page/chunk
        for doc in docs:
            doc.metadata.setdefault("source", path.name)
            doc.metadata["file_path"] = str(path)
            doc.metadata["file_type"] = ext

        logger.info(f"Loaded {len(docs)} page(s) from '{path.name}'")
        return docs

    def chunk_documents(self, docs: List[Document]) -> List[Document]:
        """
        Split a list of Documents into smaller, overlapping chunks.

        Args:
            docs: Raw documents from a loader.

        Returns:
            List of chunked Documents with preserved metadata.
        """
        if not docs:
            return []

        chunks = self.text_splitter.split_documents(docs)

        # Filter out empty or whitespace-only chunks
        chunks = [c for c in chunks if c.page_content.strip()]

        logger.info(
            f"Created {len(chunks)} chunk(s) from {len(docs)} document(s)."
        )
        return chunks

    def process_document(self, file_path: str) -> List[Document]:
        """
        Full pipeline: load → chunk a single document.

        Args:
            file_path: Path to the document file.

        Returns:
            List of chunked Document objects ready for embedding.
        """
        logger.info(f"Processing document: {file_path}")
        docs = self.load_document(file_path)
        return self.chunk_documents(docs)

    def process_multiple(self, file_paths: List[str]) -> List[Document]:
        """
        Process multiple documents, collecting all chunks.
        Files that fail are logged and skipped — they do not abort the batch.

        Args:
            file_paths: List of absolute file paths.

        Returns:
            Combined list of chunks from all successfully processed documents.
        """
        all_chunks: List[Document] = []

        for path in file_paths:
            try:
                chunks = self.process_document(path)
                all_chunks.extend(chunks)
            except Exception as e:
                logger.warning(f"Skipping '{path}': {e}")

        logger.info(
            f"Processed {len(file_paths)} file(s) → {len(all_chunks)} total chunks."
        )
        return all_chunks
